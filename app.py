import streamlit as st
import openai
import tempfile
import os
from PyPDF2 import PdfReader
from docx import Document
from docx import Document as DocxWriter
import requests
from bs4 import BeautifulSoup

openai.api_key = os.getenv("OPENAI_API_KEY")

def extract_pdf_text(file):
    reader = PdfReader(file)
    return "".join(page.extract_text() or "" for page in reader.pages)[:8000]

def extract_docx_text(file):
    doc = Document(file)
    return "\n".join(para.text for para in doc.paragraphs)[:8000]

def fetch_lawinsider_examples():
    url = "https://www.lawinsider.com/clause/scope-of-work"
    response = requests.get(url)
    soup = BeautifulSoup(response.text, "html.parser")
    clauses = soup.select(".clause-body")
    return [clause.get_text(strip=True) for clause in clauses[:5]]

def fetch_text_from_url(url):
    try:
        response = requests.get(url)
        soup = BeautifulSoup(response.text, "html.parser")
        paragraphs = soup.find_all('p')
        return "\n".join(p.get_text(strip=True) for p in paragraphs[:10])
    except Exception as e:
        return f"[Error fetching URL content: {e}]"

def generate_sow(base_text, user_desc, selected_examples):
    examples_text = "\n---\n".join(selected_examples) if selected_examples else "None included"
    prompt = f'''
You are a legal AI assistant. Based on the following base contract text, user description, and example SoW clauses, generate a detailed Statement of Work (SoW).

---
User Description:
{user_desc}

---
Base Document Extract:
{base_text}

---
Example SoWs:
{examples_text}

---
Generate the SoW using the following structure:

1. Description – What is being supplied or done.
2. Function – The business purpose or outcome the goods/services serve.
3. Price – Pricing structure, billing frequency, and payment terms.
4. Milestones – Key deliverables with corresponding deadlines or phases.
5. Warranties – Any performance guarantees, service warranties, or coverage periods.
6. Service Levels (if applicable) – SLAs, KPIs, uptime, penalties, or escalation paths.
7. Others – Any additional relevant clauses not captured above (e.g. assumptions, subcontracting, ownership of deliverables).

Also suggest questions for missing or unclear details.
'''
    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a contract lawyer."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.5
    )
    return response['choices'][0]['message']['content']

def export_to_docx(content):
    doc = DocxWriter()
    doc.add_heading("Generated Statement of Work", level=1)
    for paragraph in content.split('\n'):
        doc.add_paragraph(paragraph)
    temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_path.name)
    return temp_path.name

st.title("AI Statement of Work (SoW) Generator")

uploaded_file = st.file_uploader("Upload base contract or scope document (PDF or DOCX)", type=["pdf", "docx"])
user_desc = st.text_area("Describe the goods/services and business context")

use_examples = st.checkbox("Include example SoWs from LawInsider", value=True)
selected_examples = []
custom_examples_input = st.text_area("Paste your own SoW clauses or content here (optional)")
external_url = st.text_input("Paste a URL to extract external SoW-style clauses (optional)")

if use_examples:
    with st.spinner("Fetching LawInsider examples..."):
        all_examples = fetch_lawinsider_examples()
    st.markdown("**Select examples to include:**")
    selected_examples = [ex for i, ex in enumerate(all_examples) if st.checkbox(f"Include Example {i+1}", value=True)]

if st.button("Generate SoW"):
    if uploaded_file and user_desc:
        with st.spinner("Extracting text and generating SoW..."):
            text = extract_pdf_text(uploaded_file) if uploaded_file.name.endswith(".pdf") else extract_docx_text(uploaded_file)
            combined_examples = selected_examples.copy()
            if custom_examples_input.strip():
                combined_examples.append(custom_examples_input.strip())
            if external_url.strip():
                combined_examples.append(fetch_text_from_url(external_url.strip()))
            sow = generate_sow(text, user_desc, combined_examples)
            st.subheader("Generated Statement of Work")
            st.write(sow)
            docx_path = export_to_docx(sow)
            with open(docx_path, "rb") as f:
                st.download_button("Download SoW as DOCX", f, file_name="Statement_of_Work.docx")
    else:
        st.warning("Please upload a document and provide a description.")